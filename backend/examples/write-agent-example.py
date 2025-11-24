import os
import time
from typing import TypedDict, Annotated, List, Dict, Optional
from langgraph.graph import StateGraph, END
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, SystemMessage
import operator
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# ============================================
# 1. 配置 API
# ============================================

class LLMConfig:
    """LLM配置类,支持自定义API Key和API URL"""

    def __init__(
            self,
            api_key: str,
            api_url: str = None,
            model: str = "gpt-4",
            max_retries: int = 3,
            timeout: int = 60
    ):
        self.api_key = api_key
        self.api_url = api_url or "https://api.openai.com/v1"
        self.model = model
        self.max_retries = max_retries
        self.timeout = timeout

    def get_llm(self, temperature: float = 0.7):
        """获取配置好的LLM实例"""
        return ChatOpenAI(
            model=self.model,
            temperature=temperature,
            api_key=self.api_key,
            base_url=self.api_url,
            max_retries=self.max_retries,
            timeout=self.timeout
        )


# ============================================
# 2. 定义状态
# ============================================

class DocumentState(TypedDict):
    """文档生成状态"""
    user_request: str
    analysis: str
    assigned_expert: str
    outline: List[Dict[str, str]]
    section_assignments: Dict[str, str]
    completed_sections: Annotated[List[Dict[str, str]], operator.add]
    section_summaries: Annotated[List[Dict[str, str]], operator.add]  # 新增：章节概述列表
    final_document: str
    llm_config: LLMConfig
    error_message: Optional[str]


# ============================================
# 3. 定义专家类型
# ============================================

EXPERT_TYPES = {
    "computer_science": {
        "name": "计算机科学专家",
        "description": "擅长软件开发、算法、系统架构、人工智能等技术领域"
    },
    "legal_contract": {
        "name": "法律合同专家",
        "description": "擅长合同起草、法律条款、合规审查、知识产权等法律事务"
    },
    "business_plan": {
        "name": "商业计划专家",
        "description": "擅长商业策略、市场分析、财务规划、商业模式设计"
    },
    "medical_health": {
        "name": "医疗健康专家",
        "description": "擅长医学知识、健康管理、医疗政策、临床研究"
    },
    "finance_investment": {
        "name": "金融投资专家",
        "description": "擅长投资分析、金融产品、风险管理、资产配置"
    },
    "marketing_branding": {
        "name": "市场营销专家",
        "description": "擅长品牌策划、市场推广、用户增长、内容营销"
    },
    "education_training": {
        "name": "教育培训专家",
        "description": "擅长课程设计、教学方法、培训体系、教育技术"
    },
    "research_academic": {
        "name": "学术研究专家",
        "description": "擅长学术写作、研究方法、文献综述、论文撰写"
    }
}


# ============================================
# 4. 辅助函数
# ============================================

def safe_llm_invoke(llm, messages, retry_count=3, delay=2):
    """安全的LLM调用,带重试机制"""
    for attempt in range(retry_count):
        try:
            response = llm.invoke(messages)
            return response
        except Exception as e:
            print(f"⚠️  API调用失败 (尝试 {attempt + 1}/{retry_count}): {str(e)}")
            if attempt < retry_count - 1:
                wait_time = delay * (attempt + 1)
                print(f"   等待 {wait_time} 秒后重试...")
                time.sleep(wait_time)
            else:
                raise Exception(f"API调用失败,已重试{retry_count}次: {str(e)}")


# ============================================
# 5. 工作流节点函数
# ============================================

def analyze_request(state: DocumentState) -> DocumentState:
    """分析专家:分析用户请求并指派合适的领域专家"""
    try:
        llm = state["llm_config"].get_llm(temperature=0.3)

        expert_list = "\n".join([
            f"- {key}: {info['name']} - {info['description']}"
            for key, info in EXPERT_TYPES.items()
        ])

        prompt = f"""你是一位分析专家,需要分析用户的文档需求并指派合适的领域专家。

        可选的专家类型:
        {expert_list}

        用户请求:{state['user_request']}

        请完成以下任务:
        1. 分析用户请求的核心主题和目标
        2. 识别涉及的主要领域
        3. 选择最合适的一位主要专家来负责这个文档
        4. 简要说明为什么选择这位专家

        请按以下格式输出:
        分析:[你的分析]
        指派专家:[专家类型的key,如computer_science]
        理由:[选择理由]"""

        messages = [SystemMessage(content=prompt)]
        response = safe_llm_invoke(llm, messages)

        analysis_text = response.content

        assigned_expert = "computer_science"
        for line in analysis_text.split("\n"):
            if "指派专家:" in line or "指派专家:" in line:
                expert_key = line.split(":")[-1].split(":")[-1].strip()
                if expert_key in EXPERT_TYPES:
                    assigned_expert = expert_key
                    break

        state["analysis"] = analysis_text
        state["assigned_expert"] = assigned_expert

        print(f"\n{'=' * 60}")
        print(f"📊 分析阶段完成")
        print(f"{'=' * 60}")
        print(f"分析结果:\n{analysis_text}\n")

        return state

    except Exception as e:
        print(f"❌ 分析阶段出错: {str(e)}")
        state["error_message"] = f"分析阶段错误: {str(e)}"
        raise


def create_outline(state: DocumentState) -> DocumentState:
    """主要专家:创建文档大纲"""
    try:
        llm = state["llm_config"].get_llm(temperature=0.5)

        expert_info = EXPERT_TYPES[state["assigned_expert"]]

        prompt = f"""你是一位{expert_info['name']},{expert_info['description']}。

        用户请求:{state['user_request']}

        分析结果:{state['analysis']}

        请为这份文档创建详细的大纲。大纲应该包含2-3个主要章节,每个章节都有明确的主题。

        请按以下格式输出每个章节(每行一个章节):
        第1章:[章节标题] | [简要描述] | [建议专家类型]
        第2章:[章节标题] | [简要描述] | [建议专家类型]
        ...

        建议专家类型请从以下选择:{', '.join(EXPERT_TYPES.keys())}"""

        messages = [SystemMessage(content=prompt)]
        response = safe_llm_invoke(llm, messages)

        outline = []
        for line in response.content.split("\n"):
            if line.strip() and "第" in line and "章" in line:
                parts = line.split("|")
                if len(parts) >= 3:
                    title = parts[0].strip()
                    description = parts[1].strip()
                    expert = parts[2].strip()

                    expert_key = expert
                    for key in EXPERT_TYPES.keys():
                        if key in expert.lower():
                            expert_key = key
                            break

                    outline.append({
                        "title": title,
                        "description": description,
                        "expert": expert_key
                    })

        if not outline:
            outline = [
                {
                    "title": f"第1章:{state['user_request']}概述",
                    "description": "文档主题的整体介绍",
                    "expert": state["assigned_expert"]
                },
                {
                    "title": "第2章:核心内容",
                    "description": "详细阐述主要内容",
                    "expert": state["assigned_expert"]
                },
                {
                    "title": "第3章:总结与展望",
                    "description": "总结要点并展望未来",
                    "expert": state["assigned_expert"]
                }
            ]

        state["outline"] = outline

        print(f"\n{'=' * 60}")
        print(f"📝 大纲创建完成")
        print(f"{'=' * 60}")
        for item in outline:
            print(f"{item['title']}")
            print(f"  描述: {item['description']}")
            print(f"  专家: {EXPERT_TYPES[item['expert']]['name']}\n")

        return state

    except Exception as e:
        print(f"❌ 大纲创建出错: {str(e)}")
        state["error_message"] = f"大纲创建错误: {str(e)}"
        raise


def assign_sections(state: DocumentState) -> DocumentState:
    """分配章节给不同的专家"""
    assignments = {}
    for i, section in enumerate(state["outline"]):
        section_id = f"section_{i + 1}"
        assignments[section_id] = section["expert"]

    state["section_assignments"] = assignments
    if "completed_sections" not in state:
        state["completed_sections"] = []
    if "section_summaries" not in state:
        state["section_summaries"] = []

    return state


def write_section(state: DocumentState) -> DocumentState:
    """领域专家:编写具体章节"""
    try:
        completed_count = len(state.get("completed_sections", []))

        if completed_count >= len(state["outline"]):
            return state

        section = state["outline"][completed_count]
        expert_type = section["expert"]
        expert_info = EXPERT_TYPES[expert_type]

        llm = state["llm_config"].get_llm(temperature=0.7)

        # 构建之前章节的概述上下文
        previous_summaries = ""
        if state.get("section_summaries"):
            previous_summaries = "\n\n已完成章节的概述:\n"
            for i, summary in enumerate(state["section_summaries"]):
                previous_summaries += f"\n{summary['title']}\n概述: {summary['summary']}\n"

        prompt = f"""你是一位{expert_info['name']},{expert_info['description']}。

整体文档主题:{state['user_request']}

整体分析:{state['analysis']}
{previous_summaries}

你需要编写的章节:
{section['title']}
章节描述:{section['description']}

请编写这一章节的详细内容,要求:
1. 内容专业、详实、有深度
2. 逻辑清晰,结构合理
3. 字数在800-1500字之间
4. 包含具体的案例、数据或方法(如适用)
5. 与之前的章节保持连贯性和一致性

请直接输出章节内容,不需要重复章节标题。"""

        messages = [SystemMessage(content=prompt)]
        response = safe_llm_invoke(llm, messages)

        completed_section = {
            "title": section["title"],
            "content": response.content,
            "expert": expert_type
        }

        print(f"\n{'=' * 60}")
        print(f"✍️  章节编写完成 ({completed_count + 1}/{len(state['outline'])}): {section['title']}")
        print(f"   专家: {expert_info['name']}")
        print(f"   字数: {len(response.content)}")
        print(f"{'=' * 60}\n")

        # 生成章节概述
        summary = generate_section_summary(
            llm=llm,
            section_title=section["title"],
            section_content=response.content,
            expert_name=expert_info['name']
        )

        section_summary = {
            "title": section["title"],
            "summary": summary
        }

        print(f"📋 章节概述已生成 ({len(summary)}字)\n")
        print(f"概述内容:\n{summary}\n")

        return {
            "completed_sections": [completed_section],
            "section_summaries": [section_summary]
        }

    except Exception as e:
        print(f"❌ 章节编写出错: {str(e)}")
        state["error_message"] = f"章节编写错误: {str(e)}"
        raise


def generate_section_summary(llm, section_title: str, section_content: str, expert_name: str) -> str:
    """
    生成章节的300字概述

    参数:
        llm: LLM实例
        section_title: 章节标题
        section_content: 章节内容
        expert_name: 专家名称

    返回:
        300字左右的章节概述
    """
    prompt = f"""你是一位{expert_name},刚刚完成了以下章节的编写。

章节标题:{section_title}

章节内容:
{section_content}

请为这个章节编写一个300字左右的概述,要求:
1. 准确概括章节的核心内容和要点
2. 突出关键信息和主要观点
3. 保持专业性和准确性
4. 字数控制在280-320字之间
5. 直接输出概述内容,不需要任何前缀或标题

请开始编写概述:"""

    messages = [SystemMessage(content=prompt)]
    response = safe_llm_invoke(llm, messages)

    return response.content.strip()


def should_continue_writing(state: DocumentState) -> str:
    """判断是否继续编写章节"""
    completed_count = len(state.get("completed_sections", []))
    total_sections = len(state["outline"])

    if completed_count < total_sections:
        return "continue"
    else:
        return "done"


def compile_document(state: DocumentState) -> DocumentState:
    """直接拼接所有章节,生成最终文档(不调用LLM)"""
    try:
        print(f"\n{'=' * 60}")
        print(f"📦 开始组装文档...")
        print(f"{'=' * 60}\n")

        # 组装文档标题
        document_parts = [f"# {state['user_request']}\n\n"]

        # 按顺序拼接所有章节
        for section in state["completed_sections"]:
            document_parts.append(f"## {section['title']}\n\n")
            document_parts.append(f"{section['content']}\n\n")

        # 生成最终文档
        final_document = "".join(document_parts)
        state["final_document"] = final_document

        print(f"✅ 文档组装完成!")
        print(f"   总字数: {len(final_document)}")
        print(f"   章节数: {len(state['completed_sections'])}")
        print(f"   概述数: {len(state.get('section_summaries', []))}\n")

        return state

    except Exception as e:
        print(f"❌ 文档组装出错: {str(e)}")
        state["error_message"] = f"文档组装错误: {str(e)}"
        raise


def save_to_word(content: str, filename: str, title: str = None, summaries: List[Dict[str, str]] = None):
    """
    将文档内容保存为Word格式

    参数:
        content: 文档内容(Markdown格式)
        filename: 输出文件名
        title: 文档标题
        summaries: 章节概述列表(可选)
    """
    doc = Document()

    # 1. 设置文档标题
    if title:
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()  # 添加空行

    # 2. 添加目录
    doc.add_heading("目录", level=1)

    # 解析内容获取章节信息用于生成目录
    lines = content.split('\n')
    chapter_num = 1
    for line in lines:
        line = line.strip()
        if line.startswith('## '):
            chapter_title = line[3:].strip()
            doc.add_paragraph(f"{chapter_num}. {chapter_title}", style='List Number')
            chapter_num += 1

    doc.add_page_break()  # 目录后分页

    # 3. 解析并添加正文内容
    for line in lines:
        line = line.strip()

        if not line:
            continue

        # 处理一级标题(# 标题) - 跳过,因为已作为文档标题
        if line.startswith('# '):
            continue

        # 处理二级标题(## 标题)
        elif line.startswith('## '):
            text = line[3:].strip()
            doc.add_heading(text, level=2)

        # 处理三级标题(### 标题)
        elif line.startswith('### '):
            text = line[4:].strip()
            doc.add_heading(text, level=3)

        # 处理普通段落
        else:
            text = line.replace('**', '')
            if text:
                paragraph = doc.add_paragraph(text)
                paragraph_format = paragraph.paragraph_format
                paragraph_format.line_spacing = 1.5
                paragraph_format.space_after = Pt(6)

    # 保存文档
    doc.save(filename)
    print(f"\n{'=' * 60}")
    print(f"📄 Word文档已保存到: {filename}")
    print(f"{'=' * 60}\n")


# ============================================
# 6. 构建工作流图
# ============================================

def create_document_workflow():
    """创建文档生成工作流"""
    workflow = StateGraph(DocumentState)

    workflow.add_node("analyze", analyze_request)
    workflow.add_node("create_outline", create_outline)
    workflow.add_node("assign_sections", assign_sections)
    workflow.add_node("write_section", write_section)
    workflow.add_node("compile", compile_document)

    workflow.set_entry_point("analyze")
    workflow.add_edge("analyze", "create_outline")
    workflow.add_edge("create_outline", "assign_sections")
    workflow.add_edge("assign_sections", "write_section")

    workflow.add_conditional_edges(
        "write_section",
        should_continue_writing,
        {
            "continue": "write_section",
            "done": "compile"
        }
    )

    workflow.add_edge("compile", END)

    return workflow.compile()


# ============================================
# 7. 主函数
# ============================================

def generate_document(
        user_request: str,
        api_key: str,
        api_url: str = "https://api.openai.com/v1",
        model: str = "gpt-4",
        max_retries: int = 3,
        timeout: int = 120,
        output_format: str = "docx"  # 输出格式
) -> str:
    """
    生成长文档

    参数:
        user_request: 用户的文档需求
        api_key: OpenAI API Key
        api_url: API 基础URL(默认OpenAI官方)
        model: 使用的模型(默认gpt-4)
        max_retries: 最大重试次数
        timeout: 超时时间(秒)
        output_format: 输出格式,"docx" 或 "md"

    返回:
        生成的完整文档
    """
    if not api_key:
        raise ValueError("API Key 不能为空")

    llm_config = LLMConfig(
        api_key=api_key,
        api_url=api_url,
        model=model,
        max_retries=max_retries,
        timeout=timeout
    )

    app = create_document_workflow()

    initial_state = {
        "user_request": user_request,
        "llm_config": llm_config,
        "completed_sections": [],
        "section_summaries": [],
        "error_message": None
    }

    print(f"\n{'=' * 60}")
    print(f"🚀 开始生成文档")
    print(f"{'=' * 60}")
    print(f"用户需求: {user_request}")
    print(f"API URL: {api_url}")
    print(f"模型: {model}")
    print(f"输出格式: {output_format.upper()}\n")

    try:
        final_state = app.invoke(initial_state)

        if final_state.get("error_message"):
            print(f"\n⚠️  警告: {final_state['error_message']}")

        document = final_state.get("final_document", "文档生成失败")
        summaries = final_state.get("section_summaries", [])

        # 保存文档
        if output_format.lower() == "docx":
            output_file = "generated_document.docx"
            save_to_word(document, output_file, user_request, summaries)
        else:
            output_file = "generated_document.md"
            with open(output_file, "w", encoding="utf-8") as f:
                # 1. 添加标题
                f.write(f"# {user_request}\n\n")

                # 2. 添加目录
                f.write("## 目录\n\n")
                lines = document.split('\n')
                chapter_num = 1
                for line in lines:
                    line = line.strip()
                    if line.startswith('## '):
                        chapter_title = line[3:].strip()
                        f.write(f"{chapter_num}. {chapter_title}\n")
                        chapter_num += 1
                f.write("\n---\n\n")

                # 3. 添加正文（去除原有的一级标题）
                for line in lines:
                    if not line.strip().startswith('# '):
                        f.write(line + '\n')

        return document

    except Exception as e:
        print(f"\n❌ 文档生成失败: {str(e)}")
        print("\n请检查:")
        print("1. API Key 是否正确")
        print("2. API URL 是否可访问")
        print("3. 网络连接是否正常")
        print("4. 是否需要配置代理")
        raise


# ============================================
# 8. 使用示例
# ============================================

if __name__ == "__main__":
    API_KEY = "sk-PagpyNwAitZjhGzzrWKUssACvvFk28U8O4dmypPxEkrRo2jh"
    API_URL = "https://xiaoai.plus/v1"

    user_request = "编写一份关于区块链技术在供应链管理中应用的研究报告"

    try:
        document = generate_document(
            user_request=user_request,
            api_key=API_KEY,
            api_url=API_URL,
            model="gpt-4o-mini",
            max_retries=3,
            timeout=120,
            output_format="docx"  # 输出为Word文档
        )

        print("文档预览(前500字):\n")
        print(document[:500] + "...\n")

    except Exception as e:
        print(f"\n程序执行失败: {str(e)}")